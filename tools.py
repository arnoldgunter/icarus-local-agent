import os
import json
import csv
import mimetypes
import hashlib
from pathlib import Path
from datetime import datetime
import requests
from bs4 import BeautifulSoup
from dotenv import load_dotenv

load_dotenv()

FIRECRAWL_API_KEY = os.getenv("FIRECRAWL_API_KEY")

ALLOWED_ROOTS = [
    Path.home().resolve(),
    (Path.home() / "icarus_uploads").resolve()
]

MAX_TEXT_CHARS = 60_000
MAX_FILE_BYTES = 100 * 1024 * 1024
MAX_DIR_ITEMS = 500

IMAGE_EXTENSIONS = {
    ".png",
    ".jpg",
    ".jpeg",
    ".webp",
    ".gif",
    ".bmp",
    ".tiff",
    ".tif"
}


def normalize_path(path):
    if not isinstance(path, str) or not path.strip():
        raise ValueError("Missing path.")

    return Path(path).expanduser().resolve()


def is_path_inside_allowed_roots(path):
    resolved = normalize_path(str(path))

    for root in ALLOWED_ROOTS:
        try:
            resolved.relative_to(root)
            return True
        except ValueError:
            continue

    return False


def assert_safe_read_path(path):
    resolved = normalize_path(path)

    if not is_path_inside_allowed_roots(resolved):
        raise PermissionError(f"Access denied outside allowed roots: {resolved}")

    if not resolved.exists():
        raise FileNotFoundError(f"Path does not exist: {resolved}")

    return resolved


def assert_safe_file(path):
    resolved = assert_safe_read_path(path)

    if not resolved.is_file():
        raise ValueError(f"Not a file: {resolved}")

    size = resolved.stat().st_size

    if size > MAX_FILE_BYTES:
        raise ValueError(f"File too large: {size} bytes. Limit: {MAX_FILE_BYTES} bytes.")

    return resolved


def is_image_file(path):
    try:
        resolved = Path(path)
        suffix = resolved.suffix.lower()

        if suffix in IMAGE_EXTENSIONS:
            return True

        mime, _ = mimetypes.guess_type(str(resolved))
        return bool(mime and mime.startswith("image/"))

    except Exception:
        return False


def truncate_text(text, max_chars=MAX_TEXT_CHARS):
    if text is None:
        return ""

    text = str(text)

    if len(text) <= max_chars:
        return text

    return text[:max_chars] + f"\n\n[TRUNCATED after {max_chars} characters]"


def basic_file_info(path):
    resolved = assert_safe_read_path(path)
    stat = resolved.stat()

    mime, encoding = mimetypes.guess_type(str(resolved))

    return {
        "path": str(resolved),
        "name": resolved.name,
        "suffix": resolved.suffix.lower(),
        "is_file": resolved.is_file(),
        "is_dir": resolved.is_dir(),
        "size_bytes": stat.st_size,
        "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(timespec="seconds"),
        "mime": mime,
        "encoding_guess": encoding,
        "is_image": is_image_file(str(resolved))
    }


def list_available_tools():
    return {
        "file_info": {
            "description": "Return metadata for a file or directory.",
            "args": {"path": "absolute path"}
        },
        "list_directory": {
            "description": "List files and folders inside a directory.",
            "args": {"path": "absolute path", "recursive": "optional bool"}
        },
        "read_text_file": {
            "description": "Read plain text, markdown, code, logs and similar text files.",
            "args": {"path": "absolute path"}
        },
        "read_pdf": {
            "description": "Extract text from a PDF file.",
            "args": {"path": "absolute path"}
        },
        "read_docx": {
            "description": "Extract text and tables from a Word .docx document.",
            "args": {"path": "absolute path"}
        },
        "read_image_metadata": {
            "description": "Return image metadata only. Actual visual understanding is handled by sending images directly to Ollama.",
            "args": {"path": "absolute path"}
        },
        "read_csv": {
            "description": "Read a CSV file preview and basic stats.",
            "args": {"path": "absolute path", "delimiter": "optional string"}
        },
        "read_json": {
            "description": "Read and parse a JSON file.",
            "args": {"path": "absolute path"}
        },
        "search_text": {
            "description": "Search for a string or regex inside a text-readable file.",
            "args": {"path": "absolute path", "query": "string", "regex": "optional bool"}
        },
        "hash_file": {
            "description": "Calculate sha256 hash of a file.",
            "args": {"path": "absolute path"}
        },
        "fetch_url": {
            "description": "Fetch and extract readable text from a webpage.",
            "args": {"url": "https://example.com"}
        },
        "firecrawl_search": {
            "description": "Search the web with Firecrawl and optionally scrape markdown content from results.",
            "args": {
                "query": "search query",
                "limit": "optional number, default 3",
                "scrape": "optional bool, default true"
            }
        },
        "firecrawl_fetch_url": {
            "description": "Fetch a single webpage with Firecrawl and return clean markdown.",
            "args": {
                "url": "https://example.com"
            }
        }
    }


def run_tool(tool_name, args):
    if not isinstance(args, dict):
        args = {}

    try:
        if tool_name == "file_info":
            return tool_file_info(args)

        if tool_name == "list_directory":
            return tool_list_directory(args)

        if tool_name == "read_text_file":
            return tool_read_text_file(args)

        if tool_name == "read_pdf":
            return tool_read_pdf(args)

        if tool_name == "read_docx":
            return tool_read_docx(args)

        if tool_name == "read_image_metadata":
            return tool_read_image_metadata(args)

        if tool_name == "read_csv":
            return tool_read_csv(args)

        if tool_name == "read_json":
            return tool_read_json(args)

        if tool_name == "search_text":
            return tool_search_text(args)

        if tool_name == "hash_file":
            return tool_hash_file(args)
        
        if tool_name == "fetch_url":
            return tool_fetch_url(args)
        
        if tool_name == "firecrawl_search":
            return tool_firecrawl_search(args)

        if tool_name == "firecrawl_fetch_url":
            return tool_firecrawl_fetch_url(args)

        return {
            "ok": False,
            "error": f"Unknown tool: {tool_name}",
            "available_tools": list_available_tools()
        }

    except Exception as e:
        return {
            "ok": False,
            "tool": tool_name,
            "error": str(e)
        }


def tool_file_info(args):
    path = args.get("path")

    return {
        "ok": True,
        "tool": "file_info",
        "file": basic_file_info(path)
    }


def tool_list_directory(args):
    path = args.get("path")
    recursive = bool(args.get("recursive", False))

    resolved = assert_safe_read_path(path)

    if not resolved.is_dir():
        raise ValueError(f"Not a directory: {resolved}")

    items = []
    iterator = resolved.rglob("*") if recursive else resolved.iterdir()

    for item in iterator:
        if len(items) >= MAX_DIR_ITEMS:
            break

        try:
            stat = item.stat()
            mime, _ = mimetypes.guess_type(str(item))

            items.append({
                "name": item.name,
                "path": str(item),
                "relative_path": str(item.relative_to(resolved)),
                "is_file": item.is_file(),
                "is_dir": item.is_dir(),
                "size_bytes": stat.st_size if item.is_file() else None,
                "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(timespec="seconds"),
                "suffix": item.suffix.lower(),
                "mime": mime,
                "is_image": is_image_file(str(item))
            })
        except Exception:
            continue

    return {
        "ok": True,
        "tool": "list_directory",
        "path": str(resolved),
        "recursive": recursive,
        "count": len(items),
        "truncated": len(items) >= MAX_DIR_ITEMS,
        "items": items
    }


def tool_read_text_file(args):
    path = args.get("path")
    resolved = assert_safe_file(path)

    encodings = ["utf-8", "utf-8-sig", "latin-1"]
    last_error = None

    for enc in encodings:
        try:
            with open(resolved, "r", encoding=enc, errors="replace") as f:
                content = f.read()

            return {
                "ok": True,
                "tool": "read_text_file",
                "file": basic_file_info(str(resolved)),
                "encoding": enc,
                "content": truncate_text(content)
            }

        except Exception as e:
            last_error = e

    raise RuntimeError(f"Could not read text file: {last_error}")


def tool_read_pdf(args):
    path = args.get("path")
    resolved = assert_safe_file(path)

    try:
        from pypdf import PdfReader
    except ImportError:
        return {
            "ok": False,
            "tool": "read_pdf",
            "error": "Missing dependency: pip install pypdf"
        }

    reader = PdfReader(str(resolved))
    pages = []
    full_text = ""

    for i, page in enumerate(reader.pages):
        try:
            page_text = page.extract_text() or ""
        except Exception as e:
            page_text = f"[Error extracting page {i + 1}: {e}]"

        pages.append({
            "page": i + 1,
            "text": truncate_text(page_text, 12000)
        })

        full_text += f"\n\n--- PAGE {i + 1} ---\n\n{page_text}"

    metadata = {}

    try:
        if reader.metadata:
            metadata = {str(k): str(v) for k, v in reader.metadata.items()}
    except Exception:
        metadata = {}

    return {
        "ok": True,
        "tool": "read_pdf",
        "file": basic_file_info(str(resolved)),
        "pages_count": len(reader.pages),
        "metadata": metadata,
        "content": truncate_text(full_text),
        "pages_preview": pages[:10],
        "pages_preview_truncated": len(pages) > 10
    }


def tool_read_docx(args):
    path = args.get("path")
    resolved = assert_safe_file(path)

    if resolved.suffix.lower() != ".docx":
        raise ValueError("Only .docx files are supported. Old .doc files are not supported.")

    try:
        import docx
    except ImportError:
        return {
            "ok": False,
            "tool": "read_docx",
            "error": "Missing dependency: pip install python-docx"
        }

    document = docx.Document(str(resolved))

    paragraphs = []
    for p in document.paragraphs:
        text = p.text.strip()
        if text:
            paragraphs.append(text)

    tables = []

    for table_index, table in enumerate(document.tables):
        rows = []

        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])

        tables.append({
            "table": table_index + 1,
            "rows": rows[:50],
            "truncated": len(rows) > 50
        })

    content = "\n\n".join(paragraphs)

    if tables:
        content += "\n\n[TABLES]\n"
        for table in tables:
            content += f"\nTable {table['table']}:\n"
            for row in table["rows"]:
                content += " | ".join(row) + "\n"

    return {
        "ok": True,
        "tool": "read_docx",
        "file": basic_file_info(str(resolved)),
        "paragraph_count": len(paragraphs),
        "table_count": len(tables),
        "content": truncate_text(content),
        "tables_preview": tables[:5]
    }


def tool_read_image_metadata(args):
    path = args.get("path")
    resolved = assert_safe_file(path)

    try:
        from PIL import Image, ExifTags
    except ImportError:
        return {
            "ok": False,
            "tool": "read_image_metadata",
            "error": "Missing dependency: pip install pillow"
        }

    image = Image.open(str(resolved))
    exif_data = {}

    try:
        raw_exif = image.getexif()
        tag_map = ExifTags.TAGS

        for key, value in raw_exif.items():
            name = tag_map.get(key, str(key))
            exif_data[name] = str(value)
    except Exception:
        exif_data = {}

    return {
        "ok": True,
        "tool": "read_image_metadata",
        "file": basic_file_info(str(resolved)),
        "format": image.format,
        "mode": image.mode,
        "width": image.width,
        "height": image.height,
        "exif": exif_data,
        "note": "Visual content is not extracted here. Images should be sent directly to a multimodal Ollama model via the images field."
    }


def tool_read_csv(args):
    path = args.get("path")
    delimiter = args.get("delimiter")

    resolved = assert_safe_file(path)
    rows = []

    with open(resolved, "r", encoding="utf-8-sig", errors="replace", newline="") as f:
        sample = f.read(4096)
        f.seek(0)

        if not delimiter:
            try:
                dialect = csv.Sniffer().sniff(sample)
                delimiter = dialect.delimiter
            except Exception:
                delimiter = ","

        reader = csv.DictReader(f, delimiter=delimiter)

        for i, row in enumerate(reader):
            if i >= 100:
                break
            rows.append(dict(row))

        fieldnames = reader.fieldnames or []

    return {
        "ok": True,
        "tool": "read_csv",
        "file": basic_file_info(str(resolved)),
        "delimiter": delimiter,
        "columns": fieldnames,
        "preview_rows_count": len(rows),
        "preview_rows": rows
    }


def tool_read_json(args):
    path = args.get("path")
    resolved = assert_safe_file(path)

    with open(resolved, "r", encoding="utf-8", errors="replace") as f:
        data = json.load(f)

    preview = json.dumps(data, ensure_ascii=False, indent=2)

    return {
        "ok": True,
        "tool": "read_json",
        "file": basic_file_info(str(resolved)),
        "type": type(data).__name__,
        "content": truncate_text(preview)
    }


def tool_search_text(args):
    import re

    path = args.get("path")
    query = args.get("query")
    use_regex = bool(args.get("regex", False))

    if not query:
        raise ValueError("Missing query.")

    file_result = tool_read_text_file({"path": path})
    content = file_result["content"]
    lines = content.splitlines()
    matches = []

    if use_regex:
        pattern = re.compile(query, re.IGNORECASE)

        for idx, line in enumerate(lines, start=1):
            if pattern.search(line):
                matches.append({
                    "line": idx,
                    "text": line[:1000]
                })
    else:
        lowered = query.lower()

        for idx, line in enumerate(lines, start=1):
            if lowered in line.lower():
                matches.append({
                    "line": idx,
                    "text": line[:1000]
                })

    return {
        "ok": True,
        "tool": "search_text",
        "file": file_result["file"],
        "query": query,
        "regex": use_regex,
        "count": len(matches),
        "matches": matches[:200],
        "truncated": len(matches) > 200
    }


def tool_hash_file(args):
    path = args.get("path")
    resolved = assert_safe_file(path)

    sha256 = hashlib.sha256()

    with open(resolved, "rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            sha256.update(chunk)

    return {
        "ok": True,
        "tool": "hash_file",
        "file": basic_file_info(str(resolved)),
        "sha256": sha256.hexdigest()
    }

def tool_fetch_url(args):
    url = args.get("url")

    if not url:
        raise ValueError("Missing url")

    if not url.startswith(("http://", "https://")):
        raise ValueError("Invalid URL")

    headers = {
        "User-Agent": (
            "IcarusLocalAgent/0.1 "
            "(local personal assistant; contact: local)"
        ),
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
        "Accept-Language": "de-DE,de;q=0.9,en;q=0.8"
    }

    r = requests.get(
        url,
        headers=headers,
        timeout=15,
        allow_redirects=True
    )

    r.raise_for_status()

    content_type = r.headers.get("content-type", "").lower()

    if "text/html" not in content_type:
        return {
            "ok": False,
            "tool": "fetch_url",
            "url": url,
            "status_code": r.status_code,
            "content_type": content_type,
            "error": "Only HTML pages are supported by fetch_url."
        }

    html = r.text[:300_000]

    soup = BeautifulSoup(html, "html.parser")

    for tag in soup(["script", "style", "noscript", "svg", "canvas"]):
        tag.decompose()

    title = soup.title.get_text(" ", strip=True) if soup.title else ""

    text = soup.get_text(separator="\n")
    lines = [line.strip() for line in text.splitlines()]
    lines = [line for line in lines if line]

    clean_text = "\n".join(lines)

    return {
        "ok": True,
        "tool": "fetch_url",
        "url": url,
        "status_code": r.status_code,
        "content_type": content_type,
        "title": title,
        "content": clean_text[:60_000],
        "truncated": len(clean_text) > 60_000
    }
    
def require_firecrawl_key():
    if not FIRECRAWL_API_KEY:
        raise ValueError("Missing FIRECRAWL_API_KEY in environment.")
    return FIRECRAWL_API_KEY


def tool_firecrawl_search(args):
    api_key = require_firecrawl_key()

    query = args.get("query")
    limit = int(args.get("limit", 3))
    scrape = bool(args.get("scrape", True))

    if not query:
        raise ValueError("Missing query.")

    limit = max(1, min(limit, 10))

    payload = {
        "query": query,
        "limit": limit
    }

    if scrape:
        payload["scrapeOptions"] = {
            "formats": ["markdown", "links"]
        }

    r = requests.post(
        "https://api.firecrawl.dev/v2/search",
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        },
        json=payload,
        timeout=45
    )

    r.raise_for_status()
    data = r.json()

    return {
        "ok": True,
        "tool": "firecrawl_search",
        "query": query,
        "limit": limit,
        "scrape": scrape,
        "result": data
    }


def tool_firecrawl_fetch_url(args):
    api_key = require_firecrawl_key()

    url = args.get("url")

    if not url:
        raise ValueError("Missing url.")

    if not url.startswith(("http://", "https://")):
        raise ValueError("Invalid URL.")

    payload = {
        "url": url,
        "formats": ["markdown", "links"]
    }

    r = requests.post(
        "https://api.firecrawl.dev/v2/scrape",
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        },
        json=payload,
        timeout=45
    )

    r.raise_for_status()
    data = r.json()

    return {
        "ok": True,
        "tool": "firecrawl_fetch_url",
        "url": url,
        "result": data
    }